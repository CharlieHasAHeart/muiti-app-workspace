#!/usr/bin/env python3
import argparse
import os
import re
import sys
import warnings

warnings.filterwarnings(
    "ignore",
    message="pkg_resources is deprecated as an API.*",
    category=UserWarning,
)

from bs4 import BeautifulSoup, NavigableString
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Cm
from docxtpl import DocxTemplate
from markdown import markdown

from .env_loader import load_dotenv
from .heading_analyzer import (
    FrontMatterPlan,
    HeadingNormalizationPlan,
    build_decision_plan,
    normalize_markdown_with_decision,
)
from .template_profiles import TemplateProfile, TemplateStyleProfile, get_template_profile_by_path

load_dotenv()

_HEADING_NUM_PATTERNS = [
    r"^\s*第\s*([0-9]+|[一二三四五六七八九十百千]+)\s*(章|节|部分|篇)\s*[:：、\.\s]*",
    r"^\s*[一二三四五六七八九十百千]+\s*[、\.\)]\s*",
    r"^\s*\d+(?:\.\d+)+\s*[\.\)]?\s*",
    r"^\s*\d+\s*[、\.\)]\s*",
]


def strip_heading_number(text: str) -> str:
    if not text:
        return text
    s = text.strip()
    for pat in _HEADING_NUM_PATTERNS:
        s_new = re.sub(pat, "", s).strip()
        if s_new != s:
            s = s_new
    return s


def cleanup_placeholder_paragraph(doc, placeholder: str) -> None:
    for para in list(doc.paragraphs):
        if para.text.strip() == placeholder:
            para._element.getparent().remove(para._element)


def remove_empty_subtitle_paragraph(doc, subtitle_style_candidates: list[str]) -> None:
    candidates = set(subtitle_style_candidates)
    for para in list(doc.paragraphs):
        style_name = getattr(getattr(para, "style", None), "name", "")
        if style_name in candidates and not para.text.strip():
            para._element.getparent().remove(para._element)


def apply_style(obj, style_names: list[str]) -> bool:
    for name in style_names:
        try:
            obj.style = name
            return True
        except KeyError:
            continue
    return False


def update_fields_on_open(doc) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def next_numbering_id(numbering_root) -> int:
    num_ids = []
    for num in numbering_root.findall(qn("w:num")):
        value = num.get(qn("w:numId"))
        if value and value.isdigit():
            num_ids.append(int(value))
    return max(num_ids, default=0) + 1


def create_numbering_instance(document_part, abstract_num_id: int = 2) -> int:
    numbering_root = document_part.numbering_part.numbering_definitions._numbering
    num_id = next_numbering_id(numbering_root)
    num_xml = (
        f'<w:num xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        f'w:numId="{num_id}">'
        f'<w:abstractNumId w:val="{abstract_num_id}"/>'
        f'<w:lvlOverride w:ilvl="0"><w:startOverride w:val="1"/></w:lvlOverride>'
        f'</w:num>'
    )
    numbering_root.append(parse_xml(num_xml))
    return num_id


def apply_numbering(paragraph, num_id: int, ilvl: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.numPr
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)

    ilvl_el = num_pr.find(qn("w:ilvl"))
    if ilvl_el is None:
        ilvl_el = OxmlElement("w:ilvl")
        num_pr.append(ilvl_el)
    ilvl_el.set(qn("w:val"), str(ilvl))

    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def resolve_img_path(md_path: str, src: str) -> str:
    if os.path.isabs(src):
        return src
    base_dir = os.path.dirname(os.path.abspath(md_path))
    return os.path.join(base_dir, src)


def heading_style_names(styles: TemplateStyleProfile, level: int) -> list[str]:
    return styles.headings.get(max(1, min(level, 6)), styles.headings[1])


def add_centered_image(subdoc, img_path: str, width_cm: float, styles: TemplateStyleProfile) -> None:
    p = subdoc.add_paragraph()
    if not apply_style(p, styles.image):
        p.style = "Normal"
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))


def add_caption(subdoc, text: str, styles: TemplateStyleProfile) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    if not apply_style(p, styles.caption):
        p.style = "Normal"
    p.paragraph_format.keep_together = True


def add_heading(subdoc, text: str, style_names: list[str]) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    if not apply_style(p, style_names):
        p.style = "Heading 1"
    fmt = p.paragraph_format
    fmt.left_indent = None
    fmt.first_line_indent = None
    fmt.hanging_indent = None


def render_heading_by_level(subdoc, text: str, level: int, styles: TemplateStyleProfile) -> None:
    add_heading(subdoc, text, heading_style_names(styles, level))


def add_paragraph(subdoc, text: str, style_names: list[str]) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    if not apply_style(p, style_names):
        p.style = "Normal"


def normalize_inline_text(text: str, preserve_newlines: bool = True) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    if preserve_newlines:
        parts = [re.sub(r"[ \t]+", " ", part) for part in text.split("\n")]
        return "\n".join(parts).strip("\n")
    return re.sub(r"\s+", " ", text).strip()


def add_paragraph_with_inline_code(subdoc, p_node, style_names: list[str], inline_code_styles: list[str], preserve_newlines: bool = True, prefix_text: str = ""):
    p = subdoc.add_paragraph()
    if not apply_style(p, style_names):
        p.style = "Normal"

    if prefix_text:
        p.add_run(prefix_text)

    children = list(p_node.children)
    pending_thinspace = False
    for idx, child in enumerate(children):
        if isinstance(child, NavigableString):
            text = normalize_inline_text(str(child), preserve_newlines)
            if pending_thinspace and text.startswith(" "):
                text = "\u2009" + text[1:]
                pending_thinspace = False
            elif pending_thinspace:
                pending_thinspace = False
            if text:
                if idx + 1 < len(children) and getattr(children[idx + 1], "name", None) == "code" and text.endswith(" "):
                    text = text[:-1] + "\u2009"
                p.add_run(text)
        elif hasattr(child, "name") and child.name == "code":
            code_text = child.get_text()
            if code_text:
                run = p.add_run(code_text.replace(" ", "\u00A0"))
                apply_style(run, inline_code_styles)
            if idx + 1 < len(children) and isinstance(children[idx + 1], NavigableString):
                pending_thinspace = True
        else:
            text = child.get_text(strip=False) if hasattr(child, "get_text") else str(child)
            text = normalize_inline_text(text, preserve_newlines)
            if pending_thinspace and text.startswith(" "):
                text = "\u2009" + text[1:]
                pending_thinspace = False
            elif pending_thinspace:
                pending_thinspace = False
            if text:
                p.add_run(text)

    return p


def add_list(subdoc, list_node, ordered: bool, styles: TemplateStyleProfile, level: int = 0) -> None:
    style_names = styles.ordered_list if ordered else styles.unordered_list
    num_id = create_numbering_instance(subdoc.part, abstract_num_id=2) if ordered else None
    for li in list_node.find_all("li", recursive=False):
        li_copy = BeautifulSoup(li.encode_contents(), "html.parser")
        for nested in li_copy.find_all(["ul", "ol"]):
            nested.decompose()
        if li_copy.get_text(strip=True) or li_copy.find("code"):
            paragraph = add_paragraph_with_inline_code(
                subdoc,
                li_copy,
                style_names,
                styles.inline_code,
                preserve_newlines=False,
            )
            if ordered and num_id is not None:
                apply_numbering(paragraph, num_id, ilvl=min(level, 8))
        for nested in li.find_all(["ul", "ol"], recursive=False):
            add_list(subdoc, nested, ordered=nested.name == "ol", styles=styles, level=level + 1)


def format_language(lang: str) -> str:
    if not lang:
        return ""
    mapping = {
        "py": "Python",
        "python": "Python",
        "js": "JavaScript",
        "javascript": "JavaScript",
        "ts": "TypeScript",
        "typescript": "TypeScript",
        "json": "JSON",
        "yaml": "YAML",
        "yml": "YAML",
        "bash": "Bash",
        "sh": "Shell",
        "shell": "Shell",
    }
    return mapping.get(lang.lower(), lang.capitalize())


def add_code_block(subdoc, pre_node, styles: TemplateStyleProfile) -> None:
    code_node = pre_node.find("code")
    lang = None
    if code_node and code_node.has_attr("class"):
        for cls in code_node["class"]:
            if cls.startswith("language-"):
                lang = cls.replace("language-", "").strip()
                break

    code_text = code_node.get_text() if code_node else pre_node.get_text()
    code_text = code_text.replace("\r\n", "\n").replace("\r", "\n").rstrip("\n")

    if lang:
        p_lang = subdoc.add_paragraph(f"语言：{format_language(lang)}")
        if not apply_style(p_lang, styles.code_language):
            p_lang.style = "Normal"

    for line in code_text.split("\n"):
        p = subdoc.add_paragraph(line)
        if not apply_style(p, styles.code_block):
            p.style = "Normal"


def add_table(subdoc, table_node, styles: TemplateStyleProfile) -> None:
    rows = []
    thead = table_node.find("thead")
    if thead:
        for tr in thead.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("header", cells))
    tbody = table_node.find("tbody")
    if tbody:
        for tr in tbody.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("body", cells))
    if not thead and not tbody:
        for tr in table_node.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("body", cells))
    if not rows:
        return

    max_cols = max(len(cells) for _, cells in rows)
    table = subdoc.add_table(rows=len(rows), cols=max_cols)
    if not apply_style(table, styles.table):
        table.style = "Normal Table"

    for r_idx, (row_kind, cells) in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx in range(max_cols):
            cell = row.cells[c_idx]
            cell.text = cells[c_idx] if c_idx < len(cells) else ""
            style_names = styles.table_header_paragraph if row_kind == "header" else styles.table_body_paragraph
            for paragraph in cell.paragraphs:
                apply_style(paragraph, style_names)


def render_markdown_to_subdoc(
    subdoc,
    md_path: str,
    md_text: str,
    plan: HeadingNormalizationPlan | None = None,
    template_profile: TemplateProfile | None = None,
) -> None:
    plan = plan or HeadingNormalizationPlan()
    template_profile = template_profile or get_template_profile_by_path("backend/md2word/templates/reference.docx")
    styles = template_profile.styles
    html = markdown(md_text, extensions=["extra"])
    soup = BeautifulSoup(html, "html.parser")
    body = soup.body if soup.body else soup

    fig_index = 0
    pending_table_caption = ""
    first_h1_consumed_as_title = False

    def handle_image(src: str, caption: str) -> None:
        nonlocal fig_index
        if not src:
            return
        name = caption.strip() if caption else ""
        if not name:
            base = os.path.basename(src)
            name = os.path.splitext(base)[0]
        fig_index += 1
        caption_text = f"图{fig_index} {name}" if name else f"图{fig_index}"
        img_path = resolve_img_path(md_path, src)
        if os.path.exists(img_path):
            add_centered_image(subdoc, img_path, 15, styles)
            add_caption(subdoc, caption_text, styles)
        else:
            add_paragraph(subdoc, f"[图片缺失: {src}]", styles.paragraph)
            add_caption(subdoc, caption_text, styles)

    for node in body.children:
        if not hasattr(node, "name"):
            continue
        if node.name in {"h1", "h2", "h3", "h4", "h5", "h6"}:
            original_level = int(node.name[1])
            heading_text = strip_heading_number(node.get_text(strip=True))
            if plan.skip_first_h1_in_body and original_level == 1 and not first_h1_consumed_as_title:
                first_h1_consumed_as_title = True
                continue
            effective_level = max(1, original_level + plan.heading_shift)
            render_heading_by_level(subdoc, heading_text, effective_level, styles)
        elif node.name == "p":
            text = node.get_text(strip=True)
            if text:
                if re.match(r"^表\s*\d+\s+.+", text):
                    pending_table_caption = text
                else:
                    add_paragraph_with_inline_code(subdoc, node, styles.paragraph, styles.inline_code)
            for img in node.find_all("img"):
                handle_image(img.get("src", ""), img.get("alt", "") or "")
            for link in node.find_all("a"):
                href = link.get("href", "")
                if href.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                    handle_image(href, link.get_text(strip=True))
        elif node.name == "blockquote":
            quote_text = node.get_text("\n", strip=True)
            if quote_text:
                for line in quote_text.split("\n"):
                    stripped = line.strip()
                    if stripped.startswith("提示:") or stripped.startswith("提示："):
                        add_paragraph(subdoc, line, styles.tip_quote)
                    elif stripped.startswith("注意:") or stripped.startswith("注意："):
                        add_paragraph(subdoc, line, styles.note_quote)
                    elif stripped.startswith("警告:") or stripped.startswith("警告："):
                        add_paragraph(subdoc, line, styles.warning_quote)
                    else:
                        add_paragraph(subdoc, line, styles.quote)
        elif node.name == "img":
            handle_image(node.get("src", ""), node.get("alt", "") or "")
        elif node.name == "table":
            if pending_table_caption:
                p = subdoc.add_paragraph(pending_table_caption)
                if not apply_style(p, styles.table_caption):
                    p.style = "Normal"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pending_table_caption = ""
            add_table(subdoc, node, styles)
        elif node.name == "ul":
            add_list(subdoc, node, ordered=False, styles=styles)
        elif node.name == "ol":
            add_list(subdoc, node, ordered=True, styles=styles)
        elif node.name == "pre":
            add_code_block(subdoc, node, styles)
        elif node.name == "a":
            href = node.get("href", "")
            if href.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                handle_image(href, node.get_text(strip=True))
            else:
                add_paragraph(subdoc, node.get_text(strip=True), styles.paragraph)


def resolve_subtitle(md_text: str, front_matter_plan: FrontMatterPlan) -> str:
    if not front_matter_plan.has_subtitle:
        return ""
    return front_matter_plan.subtitle_text.strip()


def resolve_cover_title(title: str, md_text: str, plan: HeadingNormalizationPlan) -> str:
    if title.strip():
        return title.strip()
    if plan.title_text:
        return plan.title_text.strip()
    match = re.search(r"^#\s+(.+?)\s*$", md_text, flags=re.M)
    if match:
        return strip_heading_number(match.group(1).strip())
    return ""


def default_output_path(md_path: str) -> str:
    base, _ = os.path.splitext(os.path.abspath(md_path))
    return f"{base}.docx"


def convert_markdown_to_docx(
    md_path: str,
    template_path: str,
    output_path: str | None = None,
    title: str = "",
) -> str:
    md_path = os.path.abspath(md_path)
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path) if output_path else default_output_path(md_path)

    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Markdown not found: {md_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    template_profile = get_template_profile_by_path(template_path)
    decision = build_decision_plan(md_text)
    normalized_md_text = normalize_markdown_with_decision(md_text, decision)
    cover_title = resolve_cover_title(title, md_text, decision.heading_plan)
    subtitle = resolve_subtitle(normalized_md_text, decision.front_matter_plan)

    tpl = DocxTemplate(template_path)
    subdoc = tpl.new_subdoc()
    render_markdown_to_subdoc(
        subdoc,
        md_path,
        normalized_md_text,
        decision.heading_plan,
        template_profile=template_profile,
    )
    tpl.render({"main_content": subdoc, "title": cover_title, "subtitle": subtitle})
    if not subtitle:
        cleanup_placeholder_paragraph(tpl.docx, "{{subtitle}}")
        remove_empty_subtitle_paragraph(tpl.docx, template_profile.styles.subtitle)
    update_fields_on_open(tpl.docx)

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    tpl.save(output_path)
    return output_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX using a reference template.")
    parser.add_argument("-i", "--input", required=True, help="Path to the Markdown file")
    parser.add_argument(
        "-t",
        "--template",
        default="backend/md2word/templates/reference.docx",
        help="Path to the DOCX template (default: backend/md2word/templates/reference.docx)",
    )
    parser.add_argument("-o", "--output", default=None, help="Output DOCX path")
    parser.add_argument("--title", default="", help="Title used for cover/header placeholders")
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = build_parser()
    args = parser.parse_args(argv)

    try:
        out_path = convert_markdown_to_docx(
            md_path=args.input,
            template_path=args.template,
            output_path=args.output,
            title=args.title,
        )
    except Exception as exc:
        print(f"Error: {exc}", file=sys.stderr)
        return 1

    print(out_path)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
