#!/usr/bin/env python3
import argparse
import os
import re
import sys
import warnings

warnings.filterwarnings(
    "ignore",
    message="pkg_resources is deprecated as an API.*",
    category=UserWarning,
)

from docxtpl import DocxTemplate
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from markdown import markdown
from bs4 import BeautifulSoup, NavigableString

_HEADING_NUM_PATTERNS = [
    r"^\s*第\s*([0-9]+|[一二三四五六七八九十百千]+)\s*(章|节|部分|篇)\s*[:：、\.\s]*",
    r"^\s*[一二三四五六七八九十百千]+\s*[、\.\)]\s*",
    r"^\s*\d+(?:\.\d+)+\s*[\.\)]?\s*",
    r"^\s*\d+\s*[、\.\)]\s*",
]

_STYLE_ALIASES = {
    "heading 1": ["heading 1", "Heading 1", "标题 1"],
    "heading 2": ["heading 2", "Heading 2", "标题 2"],
    "heading 3": ["heading 3", "Heading 3", "标题 3"],
    "heading 4": ["heading 4", "Heading 4", "标题 4"],
    "heading 5": ["heading 5", "Heading 5", "标题 5"],
    "heading 6": ["heading 6", "Heading 6", "标题 6"],
    "heading 7": ["heading 7", "Heading 7", "标题 7"],
    "heading 8": ["heading 8", "Heading 8", "标题 8"],
    "heading 9": ["heading 9", "Heading 9", "标题 9"],
    "Normal": ["Normal", "正文"],
    "Title": ["Title", "标题"],
    "Subtitle": ["Subtitle", "副标题"],
    "Quote": ["Quote", "引用"],
    "引用块": ["引用块", "Quote", "Intense Quote"],
    "提示块": ["提示块", "引用块", "Quote"],
    "注意块": ["注意块", "引用块", "Quote"],
    "警告块": ["警告块", "引用块", "Intense Quote"],
    "Intense Quote": ["Intense Quote", "明显引用"],
    "Intense Emphasis": ["Intense Emphasis", "强调"],
    "List Paragraph": ["List Paragraph", "列表段落"],
    "No List": ["No List", "无列表"],
    "列表-无序": ["列表-无序", "List Paragraph", "List"],
    "列表-有序": ["列表-有序", "List Paragraph", "List"],
    "Default Paragraph Font": ["Default Paragraph Font", "默认段落字体"],
    "Table Grid": ["Table Grid", "表格网格"],
    "Normal Table": ["Normal Table", "普通表格"],
    "header": ["header", "页眉"],
    "footer": ["footer", "页脚"],
    "page number": ["page number", "页码"],
    "toc 1": ["toc 1", "目录 1"],
    "toc 2": ["toc 2", "目录 2"],
    "toc 3": ["toc 3", "目录 3"],
    "toc 4": ["toc 4", "目录 4"],
    "Caption": ["Caption", "题注", "图注"],
}


def strip_heading_number(text: str) -> str:
    if not text:
        return text
    s = text.strip()
    for pat in _HEADING_NUM_PATTERNS:
        s_new = re.sub(pat, "", s).strip()
        if s_new != s:
            s = s_new
    return s


def _iter_style_candidates(style_names: list[str]) -> list[str]:
    candidates = []
    seen = set()
    for name in style_names:
        for alias in _STYLE_ALIASES.get(name, [name]):
            if alias not in seen:
                candidates.append(alias)
                seen.add(alias)
    return candidates


def apply_style(obj, style_names: list[str]) -> bool:
    for name in _iter_style_candidates(style_names):
        try:
            obj.style = name
            return True
        except KeyError:
            continue
    return False


def update_fields_on_open(doc) -> None:
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def resolve_img_path(md_path: str, src: str) -> str:
    if os.path.isabs(src):
        return src
    base_dir = os.path.dirname(os.path.abspath(md_path))
    return os.path.join(base_dir, src)


def add_centered_image(subdoc, img_path: str, width_cm: float) -> None:
    p = subdoc.add_paragraph()
    if not apply_style(p, ["图片", "Normal"]):
        p.style = "Normal"
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))


def add_caption(subdoc, text: str) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    if not apply_style(p, ["图注", "Caption", "Normal"]):
        p.style = "Normal"
    p.paragraph_format.keep_together = True


def add_heading(subdoc, text: str, style_name: str) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    if not apply_style(p, [style_name, "heading 1"]):
        p.style = "Heading 1"
    fmt = p.paragraph_format
    fmt.left_indent = None
    fmt.first_line_indent = None
    fmt.hanging_indent = None


def add_paragraph(subdoc, text: str, style_name: str) -> None:
    if not text:
        return
    p = subdoc.add_paragraph(text)
    if not apply_style(p, [style_name, "Normal"]):
        p.style = "Normal"


def add_paragraph_with_inline_code(subdoc, p_node, style_name: str) -> None:
    p = subdoc.add_paragraph()
    if not apply_style(p, [style_name, "Normal"]):
        p.style = "Normal"

    children = list(p_node.children)
    pending_thinspace = False
    for idx, child in enumerate(children):
        if isinstance(child, NavigableString):
            text = str(child)
            if pending_thinspace and text.startswith(" "):
                text = "\u2009" + text[1:]
                pending_thinspace = False
            elif pending_thinspace:
                pending_thinspace = False
            if text:
                if idx + 1 < len(children) and getattr(children[idx + 1], "name", None) == "code" and text.endswith(" "):
                    text = text[:-1] + "\u2009"
                p.add_run(text)
        elif hasattr(child, "name") and child.name == "code":
            code_text = child.get_text()
            if code_text:
                run = p.add_run(code_text.replace(" ", "\u00A0"))
                apply_style(run, ["行内代码", "Inline Code"])
            if idx + 1 < len(children) and isinstance(children[idx + 1], NavigableString):
                pending_thinspace = True
        else:
            text = child.get_text(strip=False) if hasattr(child, "get_text") else str(child)
            if pending_thinspace and text.startswith(" "):
                text = "\u2009" + text[1:]
                pending_thinspace = False
            elif pending_thinspace:
                pending_thinspace = False
            if text:
                p.add_run(text)


def add_list(subdoc, list_node, ordered: bool) -> None:
    style_name = "列表-有序" if ordered else "列表-无序"
    for li in list_node.find_all("li", recursive=False):
        li_copy = BeautifulSoup(li.encode_contents(), "html.parser")
        for nested in li_copy.find_all(["ul", "ol"]):
            nested.decompose()
        if li_copy.get_text(strip=True) or li_copy.find("code"):
            add_paragraph_with_inline_code(subdoc, li_copy, style_name)
        for nested in li.find_all(["ul", "ol"], recursive=False):
            add_list(subdoc, nested, ordered=nested.name == "ol")


def format_language(lang: str) -> str:
    if not lang:
        return ""
    mapping = {
        "py": "Python",
        "python": "Python",
        "js": "JavaScript",
        "javascript": "JavaScript",
        "ts": "TypeScript",
        "typescript": "TypeScript",
        "json": "JSON",
        "yaml": "YAML",
        "yml": "YAML",
        "bash": "Bash",
        "sh": "Shell",
        "shell": "Shell",
    }
    return mapping.get(lang.lower(), lang.capitalize())


def add_code_block(subdoc, pre_node) -> None:
    code_node = pre_node.find("code")
    lang = None
    if code_node and code_node.has_attr("class"):
        for cls in code_node["class"]:
            if cls.startswith("language-"):
                lang = cls.replace("language-", "").strip()
                break

    code_text = code_node.get_text() if code_node else pre_node.get_text()
    code_text = code_text.replace("\r\n", "\n").replace("\r", "\n").rstrip("\n")

    if lang:
        p_lang = subdoc.add_paragraph(f"语言：{format_language(lang)}")
        if not apply_style(p_lang, ["代码语言标记", "代码块"]):
            p_lang.style = "Normal"

    for line in code_text.split("\n"):
        p = subdoc.add_paragraph(line)
        if not apply_style(p, ["代码块", "Normal"]):
            p.style = "Normal"


def add_table(subdoc, table_node) -> None:
    rows = []
    thead = table_node.find("thead")
    if thead:
        for tr in thead.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("header", cells))
    tbody = table_node.find("tbody")
    if tbody:
        for tr in tbody.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("body", cells))
    if not thead and not tbody:
        for tr in table_node.find_all("tr", recursive=False):
            cells = [cell.get_text(strip=True) for cell in tr.find_all(["th", "td"], recursive=False)]
            if cells:
                rows.append(("body", cells))
    if not rows:
        return

    max_cols = max(len(cells) for _, cells in rows)
    table = subdoc.add_table(rows=len(rows), cols=max_cols)
    if not apply_style(table, ["CyanScript Table", "Normal Table", "Table Grid"]):
        table.style = "Normal Table"

    for r_idx, (row_kind, cells) in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx in range(max_cols):
            cell = row.cells[c_idx]
            cell.text = cells[c_idx] if c_idx < len(cells) else ""
            style_candidates = ["表格-表头", "表格表头"] if row_kind == "header" else ["表格-正文", "表格正文"]
            for paragraph in cell.paragraphs:
                apply_style(paragraph, style_candidates)


def render_markdown_to_subdoc(subdoc, md_path: str) -> None:
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    html = markdown(md_text, extensions=["extra"])
    soup = BeautifulSoup(html, "html.parser")
    body = soup.body if soup.body else soup

    fig_index = 0
    pending_table_caption = ""

    def handle_image(src: str, caption: str) -> None:
        nonlocal fig_index
        if not src:
            return
        name = caption.strip() if caption else ""
        if not name:
            base = os.path.basename(src)
            name = os.path.splitext(base)[0]
        fig_index += 1
        caption_text = f"图{fig_index} {name}" if name else f"图{fig_index}"
        img_path = resolve_img_path(md_path, src)
        if os.path.exists(img_path):
            add_centered_image(subdoc, img_path, 15)
            add_caption(subdoc, caption_text)
        else:
            add_paragraph(subdoc, f"[图片缺失: {src}]", "Normal")
            add_caption(subdoc, caption_text)

    for node in body.children:
        if not hasattr(node, "name"):
            continue
        if node.name == "h1":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "heading 1")
        elif node.name == "h2":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "heading 2")
        elif node.name == "h3":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "heading 3")
        elif node.name == "h4":
            add_heading(subdoc, strip_heading_number(node.get_text(strip=True)), "heading 4")
        elif node.name == "p":
            text = node.get_text(strip=True)
            if text:
                if re.match(r"^表\s*\d+\s+.+", text):
                    pending_table_caption = text
                else:
                    add_paragraph_with_inline_code(subdoc, node, "Normal")
            for img in node.find_all("img"):
                handle_image(img.get("src", ""), img.get("alt", "") or "")
            for link in node.find_all("a"):
                href = link.get("href", "")
                if href.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                    handle_image(href, link.get_text(strip=True))
        elif node.name == "blockquote":
            quote_text = node.get_text("\n", strip=True)
            if quote_text:
                for line in quote_text.split("\n"):
                    stripped = line.strip()
                    if stripped.startswith("提示:") or stripped.startswith("提示："):
                        add_paragraph(subdoc, line, "提示块")
                    elif stripped.startswith("注意:") or stripped.startswith("注意："):
                        add_paragraph(subdoc, line, "注意块")
                    elif stripped.startswith("警告:") or stripped.startswith("警告："):
                        add_paragraph(subdoc, line, "警告块")
                    else:
                        add_paragraph(subdoc, line, "引用块")
        elif node.name == "img":
            handle_image(node.get("src", ""), node.get("alt", "") or "")
        elif node.name == "table":
            if pending_table_caption:
                p = subdoc.add_paragraph(pending_table_caption)
                if not apply_style(p, ["表注", "Caption", "Normal"]):
                    p.style = "Normal"
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pending_table_caption = ""
            add_table(subdoc, node)
        elif node.name == "ul":
            add_list(subdoc, node, ordered=False)
        elif node.name == "ol":
            add_list(subdoc, node, ordered=True)
        elif node.name == "pre":
            add_code_block(subdoc, node)
        elif node.name == "a":
            href = node.get("href", "")
            if href.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp")):
                handle_image(href, node.get_text(strip=True))
            else:
                add_paragraph(subdoc, node.get_text(strip=True), "Normal")


def default_output_path(md_path: str) -> str:
    base, _ = os.path.splitext(os.path.abspath(md_path))
    return f"{base}.docx"


def convert_markdown_to_docx(
    md_path: str,
    template_path: str,
    output_path: str | None = None,
    title: str = "",
) -> str:
    md_path = os.path.abspath(md_path)
    template_path = os.path.abspath(template_path)
    output_path = os.path.abspath(output_path) if output_path else default_output_path(md_path)

    if not os.path.exists(md_path):
        raise FileNotFoundError(f"Markdown not found: {md_path}")
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    tpl = DocxTemplate(template_path)
    subdoc = tpl.new_subdoc()
    render_markdown_to_subdoc(subdoc, md_path)
    tpl.render({"main_content": subdoc, "title": title})
    update_fields_on_open(tpl.docx)

    out_dir = os.path.dirname(output_path)
    if out_dir and not os.path.exists(out_dir):
        os.makedirs(out_dir, exist_ok=True)
    tpl.save(output_path)
    return output_path


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="Convert Markdown to Word (.docx) using a template")
    parser.add_argument("-i", "--input", help="Path to Markdown file")
    parser.add_argument("-t", "--template", help="Path to Word template (.docx)")
    parser.add_argument("-o", "--output", help="Path to output .docx (optional)")
    return parser


def main() -> None:
    parser = build_parser()
    args = parser.parse_args()

    md_path = args.input
    template_path = args.template
    output_path = args.output

    if not (md_path and template_path):
        md_path = input("[INPUT] Markdown file: ").strip()
        template_path = input("[INPUT] Word template (.docx): ").strip()
        output_path = input("[INPUT] Output file (.docx, blank for default): ").strip() or None

    try:
        final_path = convert_markdown_to_docx(md_path, template_path, output_path)
    except Exception as exc:
        print(f"[ERROR] {exc}")
        sys.exit(1)

    print(f"[OK] Generated: {final_path}")


if __name__ == "__main__":
    main()
